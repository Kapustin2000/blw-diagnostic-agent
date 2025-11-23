import pathlib
import json
from typing import Dict, Union, Optional
from google.adk.agents.llm_agent import ToolContext
from docx import Document  # python-docx
from ..sub_agents.doc_structure_planner.agent import DocStructure

def create_docx_from_structure(
    output_path: Optional[str] = None,
    language: Optional[str] = None,
    tool_context: Optional[ToolContext] = None,
) -> dict:
    """Create a .docx file from a DocStructure.

    Args:
        output_path:   Destination file path (can be relative). If not provided, will be retrieved from state or default to "diagnostic_report.docx".
        language:      Document language code ('uk' for Ukrainian, 'ru' for Russian, 'en' for English). If not provided, will be retrieved from state or default to "uk".
        tool_context:  ADK tool context – used to store the generated path and retrieve doc_structure from state.

    Returns:
        dict with a ``status`` message.
    """
    # Set defaults if not provided
    if output_path is None:
        output_path = "diagnostic_report.docx"
    if language is None:
        language = "uk"
    
    # Try to get output_path and language from state if available
    if tool_context is not None:
        if 'output_path' in tool_context.state:
            output_path = tool_context.state['output_path']
        if 'language' in tool_context.state:
            language = tool_context.state['language']
    
    # Try to get doc_structure from state first (most reliable source)
    doc_structure = None
    if tool_context is not None and 'doc_structure' in tool_context.state:
        state_structure = tool_context.state['doc_structure']
        # Prefer state structure if it exists and has the right format
        if isinstance(state_structure, dict) and 'sections' in state_structure:
            doc_structure = state_structure
        elif isinstance(state_structure, DocStructure):
            doc_structure = state_structure
    
    # Handle different input types
    if doc_structure is None:
        raise ValueError("doc_structure not found in state. Ensure doc_structure_planner_agent has been called first.")
    
    if isinstance(doc_structure, str):
        # Try to parse as JSON
        try:
            doc_structure = json.loads(doc_structure)
        except json.JSONDecodeError:
            raise ValueError(f"Could not parse doc_structure as JSON: {doc_structure[:100]}")
    
    if isinstance(doc_structure, dict):
        # Check if dict has 'sections' key (required for DocStructure)
        if 'sections' not in doc_structure:
            raise ValueError("doc_structure dict must have 'sections' key")
        # Convert dict to DocStructure model
        doc_structure = DocStructure(**doc_structure)
    elif not isinstance(doc_structure, DocStructure):
        raise TypeError(f"doc_structure must be DocStructure or dict, got {type(doc_structure)}")
    
    # Ensure directory exists
    pathlib.Path(output_path).parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    
    # Store language in state for reference
    if tool_context is not None:
        tool_context.state["document_language"] = language

    def add_heading(text: str, level: int):
        doc.add_heading(text, level=level)

    for sec in doc_structure.sections:
        add_heading(sec.title, level=1)
        
        # Add description if present
        if sec.description:
            doc.add_paragraph(sec.description)
        
        # Process elements directly in section (preferred)
        if sec.elements:
            for el in sec.elements:
                if el.type == "p":
                    doc.add_paragraph(el.content or "")
                elif el.type == "ul":
                    # Prefer structured list_items over markdown content
                    if el.list_items:
                        for item in el.list_items:
                            doc.add_paragraph(item, style="List Bullet")
                    elif el.content:
                        # Fallback to markdown parsing for backward compatibility
                        for line in el.content.split("\n"):
                            if line.strip():
                                doc.add_paragraph(line.strip(), style="List Bullet")
                elif el.type == "ol":
                    # Ordered list with structured list_items
                    if el.list_items:
                        for item in el.list_items:
                            doc.add_paragraph(item, style="List Number")
                    elif el.content:
                        # Fallback to markdown parsing
                        for line in el.content.split("\n"):
                            if line.strip():
                                doc.add_paragraph(line.strip(), style="List Number")
                elif el.type == "li":
                    # Single list item (backward compatibility)
                    doc.add_paragraph(el.content or "", style="List Number")
                elif el.type == "table":
                    # Prefer structured table_data over markdown content
                    if el.table_data and len(el.table_data) > 0:
                        # Use structured table data
                        num_rows = len(el.table_data)
                        num_cols = max(len(row) for row in el.table_data) if el.table_data else 0
                        if num_cols > 0:
                            table = doc.add_table(rows=num_rows, cols=num_cols)
                            table.style = 'Light Grid Accent 1'
                            
                            for i, row in enumerate(el.table_data):
                                for j, cell in enumerate(row):
                                    if j < num_cols:
                                        cell_paragraph = table.cell(i, j).paragraphs[0]
                                        cell_paragraph.clear()
                                        run = cell_paragraph.add_run(cell)
                                        if i == 0:  # Header row
                                            run.bold = True
                    elif el.content:
                        # Fallback to markdown table parser for backward compatibility
                        rows = [r.strip() for r in el.content.strip().split("\n") if r.strip()]
                        if rows:
                            parsed_rows = []
                            for row in rows:
                                cells = [c.strip() for c in row.strip("|").split("|")]
                                parsed_rows.append(cells)
                            
                            if parsed_rows:
                                table = doc.add_table(rows=len(parsed_rows), cols=len(parsed_rows[0]))
                                table.style = 'Light Grid Accent 1'
                                
                                for i, row in enumerate(parsed_rows):
                                    for j, cell in enumerate(row):
                                        if j < len(parsed_rows[0]):
                                            cell_paragraph = table.cell(i, j).paragraphs[0]
                                            cell_paragraph.clear()
                                            run = cell_paragraph.add_run(cell)
                                            if i == 0:
                                                run.bold = True
                elif el.type == "quote":
                    # Blockquote style
                    if el.quote_text:
                        paragraph = doc.add_paragraph(el.quote_text)
                        paragraph.style = 'Quote'
                    elif el.content:
                        paragraph = doc.add_paragraph(el.content)
                        paragraph.style = 'Quote'
                elif el.type in ["h1", "h2", "h3", "h4", "h5", "h6"]:
                    level = int(el.type[1])
                    add_heading(el.content or "", level=level)
                else:
                    doc.add_paragraph(el.content or "")
        
        # Process subsections if present (for backward compatibility)
        if sec.subsections:
            for sub in sec.subsections:
                add_heading(sub.title, level=2)
                for el in sub.elements:
                    if el.type == "p":
                        doc.add_paragraph(el.content or "")
                    elif el.type == "ul":
                        if el.list_items:
                            for item in el.list_items:
                                doc.add_paragraph(item, style="List Bullet")
                        elif el.content:
                            for line in el.content.split("\n"):
                                if line.strip():
                                    doc.add_paragraph(line.strip(), style="List Bullet")
                    elif el.type == "ol":
                        if el.list_items:
                            for item in el.list_items:
                                doc.add_paragraph(item, style="List Number")
                        elif el.content:
                            for line in el.content.split("\n"):
                                if line.strip():
                                    doc.add_paragraph(line.strip(), style="List Number")
                    elif el.type == "li":
                        doc.add_paragraph(el.content or "", style="List Number")
                    elif el.type == "table":
                        if el.table_data and len(el.table_data) > 0:
                            num_rows = len(el.table_data)
                            num_cols = max(len(row) for row in el.table_data) if el.table_data else 0
                            if num_cols > 0:
                                table = doc.add_table(rows=num_rows, cols=num_cols)
                                table.style = 'Light Grid Accent 1'
                                
                                for i, row in enumerate(el.table_data):
                                    for j, cell in enumerate(row):
                                        if j < num_cols:
                                            cell_paragraph = table.cell(i, j).paragraphs[0]
                                            cell_paragraph.clear()
                                            run = cell_paragraph.add_run(cell)
                                            if i == 0:
                                                run.bold = True
                        elif el.content:
                            rows = [r.strip() for r in el.content.strip().split("\n") if r.strip()]
                            if rows:
                                parsed_rows = []
                                for row in rows:
                                    cells = [c.strip() for c in row.strip("|").split("|")]
                                    parsed_rows.append(cells)
                                
                                if parsed_rows:
                                    table = doc.add_table(rows=len(parsed_rows), cols=len(parsed_rows[0]))
                                    table.style = 'Light Grid Accent 1'
                                    
                                    for i, row in enumerate(parsed_rows):
                                        for j, cell in enumerate(row):
                                            if j < len(parsed_rows[0]):
                                                cell_paragraph = table.cell(i, j).paragraphs[0]
                                                cell_paragraph.clear()
                                                run = cell_paragraph.add_run(cell)
                                                if i == 0:
                                                    run.bold = True
                    elif el.type == "quote":
                        if el.quote_text:
                            paragraph = doc.add_paragraph(el.quote_text)
                            paragraph.style = 'Quote'
                        elif el.content:
                            paragraph = doc.add_paragraph(el.content)
                            paragraph.style = 'Quote'
                    elif el.type in ["h1", "h2", "h3", "h4", "h5", "h6"]:
                        level = int(el.type[1])
                        add_heading(el.content or "", level=level)
                    else:
                        doc.add_paragraph(el.content or "")
        
        # Add conclusion if present
        if sec.conclusion:
            doc.add_paragraph(sec.conclusion)

    doc.save(output_path)
    # Store the absolute path and language in the agent state for later retrieval
    if tool_context is not None:
        tool_context.state["generated_docx"] = str(pathlib.Path(output_path).absolute())
        tool_context.state["document_language"] = language
    return {"status": f"Docx created at {output_path} (language: {language})"}
